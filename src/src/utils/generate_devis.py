import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from datetime import datetime
from src.database.database import get_dossiers, get_produits, get_options, get_connection
import tkinter as tk
from tkinter import filedialog
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

colorRed = RGBColor(255, 0, 0)  # Rouge
colorBlueText = RGBColor(2, 136, 190)  # Texte bleu
colorBlackText = RGBColor(0, 0, 0)  # Texte noir
colorGreyText = RGBColor(128, 128, 128)  # Texte gris
colorBlueBackground = "#B8CCE4"  # Fond bleu
colorBlueBorder = "#0288be"  # Bordure bleue
colorBlackBorder = "#000000"  # Bordure noire
colorMail = RGBColor(0, 112, 192)  # Couleur mail

def format_number(number):
    """
    Formate un nombre pour l'affichage dans les documents.
    Convertit les nombres en chaînes et gère les formats spéciaux.
    """
    try:
        if isinstance(number, str):
            # Si c'est déjà une chaîne, retourner telle quelle pour "Forfait" et "Ensemble"
            if number in ["Forfait", "Ensemble"]:
                return number
            # Sinon, essayer de convertir en float
            try:
                number = float(number.replace(',', '.'))
            except ValueError:
                return number

        if isinstance(number, (int, float)):
            # Formater le nombre avec 2 décimales et remplacer le point par une virgule
            return f"{number:,.2f}".replace(',', ' ').replace('.', ',')
        
        return str(number)

    except Exception as e:
        print(f"Erreur dans format_number: {e}")
        return str(number)

def set_cell_background_color(cell, color):
    """Applique une couleur de fond à une cellule."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_document_margins(document, top, bottom, left, right):
    """Définit les marges du document."""
    section = document.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def set_table_width(table, total_width):
    """Définit la largeur totale du tableau."""
    num_cols = len(table.columns)
    col_width = Inches(total_width / num_cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

def create_header(document):
    """Ajoute un en-tête au document."""
    header_table = document.add_table(rows=1, cols=2)
    set_table_width(header_table, 7.6377953)

    # Définir manuellement la largeur des colonnes
    total_width = 7.6377953
    widths = [Inches(total_width * 0.515), Inches(total_width * 0.485)]
    for idx, width in enumerate(widths):
        header_table.columns[idx].width = width

    # Contenu gauche
    cell_left = header_table.cell(0, 0)
    paragraphs = [
        ("NMG&CO EI", 20, colorBlueText, True, False),
        ("QUALITE - RAPIDITE - EFFICACITE", 10, colorBlueText, True, False),
        ("nmgiovan@gmail.com", 16, colorMail, True, True),
        ("07.69.98.25.66", 16, colorBlackText, True, False),
    ]

    for idx, (text, size, color, bold, underline) in enumerate(paragraphs):
        paragraph = cell_left.add_paragraph() if idx > 0 else cell_left.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        run.font.color.rgb = color
        run.bold = bold
        run.underline = underline
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)

    # Contenu droit
    cell_right = header_table.cell(0, 1)
    paragraphs = [
        ("DEVIS", 16, colorBlueText, True, False),
        ("Besoin d'un service... c'est simple... NMG&CO est là", 10, colorBlackText, True, False),
        ("TRAVAUX DIVERS - MAISON - JARDIN", 10, colorBlackText, True, False),
    ]

    for idx, (text, size, color, bold, underline) in enumerate(paragraphs):
        paragraph = cell_right.add_paragraph() if idx > 0 else cell_right.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        run.font.color.rgb = color
        run.bold = bold
        run.underline = underline
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(10)

def add_dossier_info(document, dossier):
    """Ajoute les informations du dossier au document."""
    
    # Ajouter un espace avant le tableau
    space = document.add_paragraph(f"")
    space.paragraph_format.space_after = Pt(0)

    # Encadré avec les informations principales
    table = document.add_table(rows=1, cols=1)
    set_table_width(table, 7.6377953)
    table.autofit = False
    table.columns[0].width = Inches(7.6377953)

    cell = table.cell(0, 0)
    cell.text = ""

    # Ajouter les styles
    lines = [
        ("DATE :", datetime.now().strftime('%d/%m/%Y')),
        ("N° DEVIS :", dossier[1]),
        ("ADRESSE CHANTIER :", dossier[2]),
        ("LIBELLÉ TRAVAUX :", dossier[3])
    ]

    for idx, (label, value) in enumerate(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        run_label = paragraph.add_run(label)
        run_value = paragraph.add_run(f" {value}")
        run_label.bold = True
        run_label.underline = True
        run_label.font.name = "Arial"
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = colorBlackText 
        run_value.font.name = "Arial"
        run_value.font.size = Pt(10)
        run_value.font.color.rgb = colorBlackText  
        paragraph.paragraph_format.space_before = Pt(0)  
        paragraph.paragraph_format.space_after = Pt(0)  

        if idx == 0:
            paragraph.paragraph_format.space_before = Pt(12)  # Espace avant la première ligne
        if idx == len(lines) - 1:
            paragraph.paragraph_format.space_after = Pt(12)  # Espace après la dernière ligne

    # Ajouter une bordure au tableau
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 2px
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), colorBlueBorder)  
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Ajouter un espace après le tableau
    space = document.add_paragraph(f"")
    space.paragraph_format.space_after = Pt(0)

    # Tableau avec les informations supplémentaires
    fields = [
        ("ADRESSE FACTURATION", str(dossier[4])),  # Convertir en string
        ("ACOMPTE DEMANDÉ", "50% à la signature"),
        ("MODALITÉS DE PAIEMENT", "À Réception de la facture"), 
        ("MOYEN DE PAIEMENT", str(dossier[5]))  # Convertir en string
    ]

    if dossier[6] == 1:  # Garantie décennale
        fields.insert(1, ("GARANTIE", "Décenale"))

    num_cols = len(fields)
    table = document.add_table(rows=2, cols=num_cols)
    set_table_width(table, 7.6377953)
    table.autofit = False
    widths = [Inches(7.6377953 / num_cols)] * num_cols
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

    # Définir la hauteur des lignes
    table.rows[0].height = Pt(40)  # 40% de la hauteur totale
    table.rows[1].height = Pt(60)  # 60% de la hauteur totale

    for col, (label, value) in enumerate(fields):
        cell_1 = table.cell(0, col)
        cell_2 = table.cell(1, col)
        cell_1.text = str(label)  # Convertir en string
        cell_2.text = str(value)  # Convertir en string

        for paragraph in cell_1.paragraphs:
            paragraph.style.font.name = "Arial"
            paragraph.style.font.size = Pt(10)
            paragraph.style.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)  
            for run in paragraph.runs:
                run.bold = True

        for paragraph in cell_2.paragraphs:
            paragraph.style.font.name = "Arial"
            paragraph.style.font.size = Pt(10)
            paragraph.style.font.color.rgb = RGBColor(0, 0, 0)  # Noir
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col != 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)  

        set_cell_background_color(cell_1, colorBlueBackground)

        # Centrer verticalement le texte
        cell_1.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_2.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter une bordure au tableau
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 2px
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), colorBlackBorder)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def description_dossier(document, dossier):
    """Ajoute la description du dossier au document."""
    document.add_paragraph("Description :")
    # Ajouter un run vide si la description est vide
    description_paragraph = document.add_paragraph()
    description_run = description_paragraph.add_run(dossier[7] if dossier[7] else "")

    # Style de paragraphe
    paragraphs = document.paragraphs[-2:]
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
    paragraphs[0].paragraph_format.space_after = Pt(5)
    paragraphs[0].paragraph_format.space_before = Pt(12)
    paragraphs[0].runs[0].font.size = Pt(13)
    paragraphs[0].runs[0].bold = True
    paragraphs[0].runs[0].underline = True
    
    # Vérifier si le run existe avant d'appliquer le style
    if description_run:
        description_run.font.size = Pt(10)
        description_paragraph.paragraph_format.space_after = Pt(12)

def add_produits_table(document, produits, dossier_id):  # Ajouter dossier_id comme paramètre
    """Ajoute un tableau des produits et options au document."""
    # Check if any product or option has a discount greater than 0
    show_remise_column = any(produit[5] > 0 for produit in produits)
    options = get_options(dossier_id)  # Maintenant dossier_id est défini
    if options:
        show_remise_column = show_remise_column or any(option[5] > 0 for option in options)

    # Determine the number of columns based on whether to show the "Remise" column
    num_cols = 5 if show_remise_column else 4
    table = document.add_table(rows=1, cols=num_cols)
    table.autofit = False

    # Define column widths
    total_width = 7.6377953
    if show_remise_column:
        widths = [Inches(total_width * 0.45), Inches(total_width * 0.1375), Inches(total_width * 0.1375), 
                 Inches(total_width * 0.1375), Inches(total_width * 0.1375)]
    else:
        widths = [Inches(total_width * 0.5625), Inches(total_width * 0.1458), 
                 Inches(total_width * 0.1458), Inches(total_width * 0.1458)]
    
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
        for cell in table.columns[idx].cells:
            cell.width = width

    # Table header
    hdr_cells = table.rows[0].cells
    headers = ['DÉSIGNATION', 'PRIX UNITAIRE', 'QUANTITÉ', 'TOTAL']
    if show_remise_column:
        headers.insert(3, 'REMISE')

    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        hdr_cells[idx].paragraphs[0].style.font.name = "Arial"
        hdr_cells[idx].paragraphs[0].style.font.size = Pt(10)
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[idx].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[idx].paragraphs[0].paragraph_format.space_after = Pt(0)
        hdr_cells[idx].paragraphs[0].runs[0].bold = True
        set_cell_background_color(hdr_cells[idx], colorBlueBackground)

    # Corps du tableau - Produits
    total_produits_sans_remise = 0  # Total des produits sans remise
    total_produits_avec_remise = 0  # Total des produits avec remise
    
    for produit in produits:
        row_cells = table.add_row().cells
        row_cells[0].text = produit[2]  # Designation
        row_cells[1].text = f'{format_number(produit[4])} €'  # Prix Unitaire
        row_cells[2].text = f'{format_number(produit[3])}' if produit[6] is None else f'{format_number(produit[3])} {produit[6]}'  # Quantité
        
        # Calculer d'abord le total sans remise
        try:
            if produit[3] and produit[3] not in ["Forfait", "Ensemble"]:
                quantite = float(str(produit[3]).replace(',', '.'))
                total_sans_remise = quantite * produit[4]
            else:
                total_sans_remise = produit[4]
        except (ValueError, TypeError):
            total_sans_remise = produit[4]
        
        total_avec_remise = total_sans_remise - produit[5]
        total_produits_sans_remise += total_sans_remise
        total_produits_avec_remise += total_avec_remise

        if show_remise_column:
            row_cells[3].text = f'{format_number(produit[5])} €'  # Remise
            row_cells[4].text = f'{format_number(total_avec_remise)} €'  # Total
        else:
            row_cells[3].text = f'{format_number(total_avec_remise)} €'  # Total

        for cell in row_cells:
            cell.paragraphs[0].style.font.name = "Arial"
            cell.paragraphs[0].style.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        # Apply left alignment to the first column
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Options
    if options:
        for option in options:
            row_cells = table.add_row().cells
            option_cell = row_cells[0].paragraphs[0]
            option_cell.text = ""  # Clear default text
            designation_run = option_cell.add_run(f"{option[2]} (option)")
            designation_run.font.color.rgb = colorBlueText
            designation_run.font.name = "Arial"
            designation_run.font.size = Pt(10)
            option_cell.paragraph_format.space_after = Pt(0)    
            option_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Autres cellules
            row_cells[1].text = f'{format_number(option[3])}' if option[6] is None else f'{format_number(option[3])} {option[6]}'
            row_cells[2].text = f'{format_number(option[4])} €'
            # Calculer le total sans remise pour les options
            try:
                if option[3] and option[3] not in ["Forfait", "Ensemble"]:
                    quantite = float(str(option[3]).replace(',', '.'))
                    total_sans_remise = quantite * option[4]
                else:
                    total_sans_remise = option[4]
            except (ValueError, TypeError):
                total_sans_remise = option[4]
                
            total_avec_remise = total_sans_remise - option[5]

            if show_remise_column:
                row_cells[3].text = f'{format_number(option[5])} €'
                row_cells[4].text = f'{format_number(total_avec_remise)} €'
            else:
                row_cells[3].text = f'{format_number(total_avec_remise)} €'

            # Style des autres cellules
            for cell in row_cells[1:]:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(cell.text)
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.color.rgb = colorBlueText
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)

    # Set row heights
    table.rows[0].height = Inches(0.35433071)  # 0.9 cm
    for row in table.rows[1:]:
        row.height = Inches(0.23622047)  # 0.6 cm

    # Add border to the table
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 2px
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), colorBlackBorder)
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Add space after the table
    space = document.add_paragraph(f"")
    space.paragraph_format.space_after = Pt(0)

    return total_produits_sans_remise, total_produits_avec_remise  # Retourner les deux totaux

def add_page_number(paragraph):
    """Ajoute la numérotation des pages au format 'Page X/Y'."""
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"

    # Champ PAGE
    run = paragraph.add_run()
    
    # Début du champ
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    # Instruction PAGE
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    run._r.append(instrText)

    # Séparateur
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    run._r.append(separate)

    # Texte par défaut
    t = OxmlElement('w:t')
    t.text = "1"
    run._r.append(t)

    # Fin du champ
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

    # Séparateur " / "
    run = paragraph.add_run(" / ")
    run.font.name = "Arial"

    # Champ NUMPAGES
    run = paragraph.add_run()
    
    # Début du champ
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar3)

    # Instruction NUMPAGES
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = ' NUMPAGES '
    run._r.append(instrText2)

    # Séparateur
    separate2 = OxmlElement('w:fldChar')
    separate2.set(qn('w:fldCharType'), 'separate')
    run._r.append(separate2)

    # Texte par défaut
    t2 = OxmlElement('w:t')
    t2.text = "1"
    run._r.append(t2)

    # Fin du champ
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar4)

def add_footer_to_last_page(document):
    """Ajoute un pied de page au bas de la dernière page du document."""
    # Ajouter un espace avant le footer
    space = document.add_paragraph(f"")
    space.paragraph_format.space_after = Pt(0)

    # Premier paragraphe
    paragraph1 = document.add_paragraph()
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph1.text = "Ce présent devis est valable pendant 1 mois à compter de la date d'émission, si acceptation après, le montant sera révisé.\nPour le valider, renvoyez-le à l'adresse suivante avec la mention\n« BON POUR ACCORD » suivi de votre signature"
    run1 = paragraph1.runs[0]
    run1.font.name = "Arial"
    run1.font.size = Pt(8)
    run1.font.color.rgb = colorGreyText

    # Deuxième paragraphe
    paragraph2 = document.add_paragraph()
    paragraph2.text = "NMG&CO EI - 1215 Route de la Fénerie 06580 Pégomas - Tél : 07.69.98.25.66 - nmgiovan@gmail.com\nN° SIRET : 843 241 399 000 16 - APE : 8121Z\nIdentité Bancaire : FR76 1460 7003 5569 4219 9717 134"
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = paragraph2.runs[0]
    run2.font.name = "Arial"
    run2.font.size = Pt(10)
    run2.font.color.rgb = colorBlackText

def save_document(document, dossier_id):
    """Ouvre une boîte de dialogue pour sauvegarder le document."""
    root = tk.Tk()
    root.withdraw()

    # Get dossier to access numero_dossier
    dossier = get_dossier(dossier_id)
    if not dossier:
        return False
        
    numero_dossier = dossier[1].replace('/', '-')  # Remplacer '/' par '-'

    document_name = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Document Word", "*.docx")],
        initialfile=f'devis_{numero_dossier}.docx',
        title="Enregistrer le devis sous..."
    )

    if document_name:
        try:
            document.save(document_name)
            print(f'Devis généré: {document_name}')
            return True
        except Exception as e:
            print(f"Erreur lors de la sauvegarde du devis: {e}")
            return False
    else:
        print("Enregistrement annulé")
        return False

def generate_devis(dossier_id):
    """Génère un devis pour le dossier spécifié."""
    document = Document()
    
    # Ajouter une section avec des en-têtes/pieds de page différents
    section = document.sections[0]
    section.different_first_page_header_footer = False
    section.footer_distance = Inches(0.5)

    # Créer le pied de page pour la section
    footer = section.footer
    
    # Ajouter la numérotation dans le pied de page
    page_number_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    page_number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_number_paragraph.paragraph_format.space_after = Pt(0)
    page_number_paragraph.paragraph_format.space_before = Pt(12)
    add_page_number(page_number_paragraph)

    dossier = get_dossier(dossier_id)
    if not isinstance(dossier, tuple):
        print(f"Erreur: Le dossier {dossier_id} n'a pas été trouvé ou n'est pas au bon format")
        return False

    produits = get_produits(dossier_id)
    if not produits:
        produits = []
    options = get_options(dossier_id)
    if not options:
        options = []  

    set_document_margins(document, top=0.5, bottom=0.5, left=0.5, right=0.5)
    create_header(document)
    add_dossier_info(document, dossier)
    description_dossier(document, dossier)
    total_produits_sans_remise, total_produits_avec_remise = add_produits_table(document, produits, dossier_id)
    
    # Calculate totals
    acompte_percentage = 50.0
    
    # Initialiser les totaux des options
    total_options_sans_remise = 0
    total_options_avec_remise = 0
    
    # Calculer séparément les totaux des options
    if options:
        for option in options:
            try:
                if option[3] and option[3] not in ["Forfait", "Ensemble"]:
                    quantite = float(str(option[3]).replace(',', '.'))
                    option_sans_remise = quantite * option[4]
                else:
                    option_sans_remise = option[4]
            except (ValueError, TypeError):
                option_sans_remise = option[4]
            
            option_avec_remise = option_sans_remise - option[5]
            total_options_sans_remise += option_sans_remise
            total_options_avec_remise += option_avec_remise

        # Calculs finaux
        sous_total = total_produits_sans_remise  # Total produits sans remises
        sous_total_avec_options = total_produits_sans_remise + total_options_sans_remise  # Total produits et options sans remises
        total_a_payer = total_produits_avec_remise  # Total produits avec remises
        total_a_payer_avec_options = total_produits_avec_remise + total_options_avec_remise  # Total produits et options avec remises
        
        # Calcul des acomptes
        acompte = total_a_payer * (acompte_percentage / 100)
        acompte_avec_options = total_a_payer_avec_options * (acompte_percentage / 100)
        
        paragraphs = [
            f"SOUS TOTAL : {format_number(sous_total)} €",
            f"TOTAL A PAYER : {format_number(total_a_payer)} €",
            f"ACOMPTE A VERSER : {format_number(acompte)} €",
            f"SOUS TOTAL AVEC OPTIONS : {format_number(sous_total_avec_options)} €",
            f"TOTAL A PAYER AVEC OPTIONS : {format_number(total_a_payer_avec_options)} €",
            f"ACOMPTE A VERSER AVEC OPTIONS : {format_number(acompte_avec_options)} €"
        ]
    else:
        # Sans options, les calculs sont plus simples
        sous_total = total_produits_sans_remise
        total_a_payer = total_produits_avec_remise
        acompte = total_a_payer * (acompte_percentage / 100)
        
        paragraphs = [
            f"SOUS TOTAL : {format_number(sous_total)} €",
            f"TOTAL A PAYER : {format_number(total_a_payer)} €",
            f"ACOMPTE A VERSER : {format_number(acompte)} €"
        ]

    # Ajout des paragraphes au document
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if "AVEC OPTIONS" in text:
            parts = text.split("AVEC OPTIONS")
            run = paragraph.add_run(parts[0])
            run.bold = True
            run = paragraph.add_run("AVEC OPTIONS")
            run.bold = True
            run.font.color.rgb = colorRed
            run = paragraph.add_run(parts[1])
            run.bold = True
        else:
            run = paragraph.add_run(text)
            run.bold = True
        paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph("TVA NON APPLICABLE - Art. 293 B du CGI")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.runs[0]
    run.font.size = Pt(7)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)

    add_footer_to_last_page(document)
    return save_document(document, dossier_id)

def get_dossier(dossier_id):
    """Récupère un dossier par son ID"""
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM dossiers WHERE id = ?', (dossier_id,))
        dossier = cursor.fetchone()
        conn.close()
        return dossier
    except Exception as e:
        print(f"Erreur lors de la récupération du dossier: {e}")
        return None

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python generate_devis.py <dossier_id>")
        sys.exit(1)

    dossier_id = int(sys.argv[1])
    if not generate_devis(dossier_id):
        sys.exit(1)
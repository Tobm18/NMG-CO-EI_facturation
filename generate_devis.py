import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from datetime import datetime
from database import get_dossiers, get_produits
import tkinter as tk
from tkinter import filedialog
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

colorRed = RGBColor(255, 0, 0)  # Rouge
colorBlueText = RGBColor(2, 136, 190)  # Texte bleu
colorBlackText = RGBColor(0, 0, 0)  # Texte noir
colorGreyText = RGBColor(128, 128, 128)  # Texte gris
colorBlueBackground = "#B8CCE4"  # Fond bleu
colorBlueBorder = "#0288be"  # Bordure bleue
colorBlackBorder = "#000000"  # Bordure noire
colorMail = RGBColor(0, 112, 192)  # Couleur mail

def set_cell_background_color(cell, color):
    """Applique une couleur de fond à une cellule."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_document_margins(document, top, bottom, left, right):
    """Définit les marges du document."""
    section = document.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def set_table_width(table, total_width):
    """Définit la largeur totale du tableau."""
    num_cols = len(table.columns)
    col_width = Inches(total_width / num_cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

def create_header(document):
    """Ajoute un en-tête au document."""
    header_table = document.add_table(rows=1, cols=2)
    set_table_width(header_table, 7.6377953)

    # Définir manuellement la largeur des colonnes
    total_width = 7.6377953
    widths = [Inches(total_width * 0.515), Inches(total_width * 0.485)]
    for idx, width in enumerate(widths):
        header_table.columns[idx].width = width

    # Contenu gauche
    cell_left = header_table.cell(0, 0)
    paragraphs = [
        ("NMG&CO EI", 20, colorBlueText, True, False),
        ("QUALITE - RAPIDITE - EFFICACITE", 10, colorBlueText, True, False),
        ("nmgiovar@gmail.com", 16, colorMail, True, True),
        ("07.69.98.25.66", 16, colorBlackText, True, False),
    ]

    for idx, (text, size, color, bold, underline) in enumerate(paragraphs):
        paragraph = cell_left.add_paragraph() if idx > 0 else cell_left.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        run.font.color.rgb = color
        run.bold = bold
        run.underline = underline
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)

    # Contenu droit
    cell_right = header_table.cell(0, 1)
    paragraphs = [
        ("DEVIS", 16, colorBlueText, True, False),
        ("Besoin d'un service... c'est simple... NMG&CO est là", 10, colorBlackText, True, False),
        ("TRAVAUX DIVERS - MAISON - JARDIN", 10, colorBlackText, True, False),
    ]

    for idx, (text, size, color, bold, underline) in enumerate(paragraphs):
        paragraph = cell_right.add_paragraph() if idx > 0 else cell_right.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        run.font.color.rgb = color
        run.bold = bold
        run.underline = underline
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(10)

def add_dossier_info(document, dossier):
    """Ajoute les informations du dossier au document."""
    
    # Ajouter un espace avant le tableau
    document.add_paragraph()

    # Encadré avec les informations principales
    table = document.add_table(rows=1, cols=1)
    set_table_width(table, 7.6377953)
    table.autofit = False
    table.columns[0].width = Inches(7.6377953)

    cell = table.cell(0, 0)
    cell.text = ""

    # Ajouter les styles
    lines = [
        ("DATE :", datetime.now().strftime('%d/%m/%Y')),
        ("N° DEVIS :", "2024/01"),
        ("ADRESSE CHANTIER :", dossier[2]),
        ("LIBELLÉ TRAVAUX :", dossier[3])
    ]

    for idx, (label, value) in enumerate(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        run_label = paragraph.add_run(label)
        run_value = paragraph.add_run(f" {value}")
        run_label.bold = True
        run_label.underline = True
        run_label.font.name = "Arial"
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = colorBlackText 
        run_value.font.name = "Arial"
        run_value.font.size = Pt(10)
        run_value.font.color.rgb = colorBlackText  
        paragraph.paragraph_format.space_before = Pt(0)  
        paragraph.paragraph_format.space_after = Pt(0)  

        if idx == 0:
            paragraph.paragraph_format.space_before = Pt(12)  # Espace avant la première ligne
        if idx == len(lines) - 1:
            paragraph.paragraph_format.space_after = Pt(12)  # Espace après la dernière ligne

    # Ajouter une bordure au tableau
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 2px
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), colorBlueBorder)  
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Ajouter un espace après le tableau
    document.add_paragraph()

    # Tableau avec les informations supplémentaires
    fields = [
        ("ADRESSE FACTURATION", dossier[4]),
        ("ACOMPTE DEMANDÉ", f"{dossier[5]}% à la signature"),
        ("MODALITÉS DE PAIEMENT", "À Réception de la facture"),  # Remplacez par la valeur appropriée
        ("MOYEN DE PAIEMENT", dossier[6])
    ]

    if dossier[7] == "1":
        fields.insert(1, ("GARANTIE", "Décenale"))

    num_cols = len(fields)
    table = document.add_table(rows=2, cols=num_cols)
    set_table_width(table, 7.6377953)
    table.autofit = False
    widths = [Inches(7.6377953 / num_cols)] * num_cols
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

    # Définir la hauteur des lignes
    table.rows[0].height = Pt(40)  # 40% de la hauteur totale
    table.rows[1].height = Pt(60)  # 60% de la hauteur totale

    for col, (label, value) in enumerate(fields):
        cell_1 = table.cell(0, col)
        cell_2 = table.cell(1, col)
        cell_1.text = label
        cell_2.text = value

        for paragraph in cell_1.paragraphs:
            paragraph.style.font.name = "Arial"
            paragraph.style.font.size = Pt(10)
            paragraph.style.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)  
            for run in paragraph.runs:
                run.bold = True

        for paragraph in cell_2.paragraphs:
            paragraph.style.font.name = "Arial"
            paragraph.style.font.size = Pt(10)
            paragraph.style.font.color.rgb = RGBColor(0, 0, 0)  # Noir
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col != 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)  

        set_cell_background_color(cell_1, colorBlueBackground)

        # Centrer verticalement le texte
        cell_1.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_2.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter une bordure au tableau
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 2px
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), colorBlackBorder)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def description_dossier(document, dossier):
    """Ajoute la description du dossier au document."""
    document.add_paragraph("Description :")
    document.add_paragraph(dossier[8])

    # Style de paragraphe
    paragraphs = document.paragraphs[-2:]
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
    paragraphs[0].paragraph_format.space_after = Pt(12)
    paragraphs[0].paragraph_format.space_before = Pt(12)
    paragraphs[0].runs[0].font.size = Pt(13)
    paragraphs[0].runs[0].bold = True
    paragraphs[0].runs[0].underline = True
    paragraphs[1].paragraph_format.space_after = Pt(12)
    paragraphs[1].runs[0].font.size = Pt(10)

def add_produits_table(document, produits):
    """Ajoute un tableau des produits au document."""
    table = document.add_table(rows=1, cols=5)
    table.autofit = False

    # Définir la largeur des colonnes
    total_width = 7.6377953
    widths = [Inches(total_width * 0.45), Inches(total_width * 0.1375), Inches(total_width * 0.1375), Inches(total_width * 0.1375), Inches(total_width * 0.1375)]
    
    table.columns[0].width = widths[0]
    table.columns[1].width = widths[1]
    table.columns[2].width = widths[2]
    table.columns[3].width = widths[3]
    table.columns[4].width = widths[4]

    for cell in table.columns[0].cells:
        cell.width = widths[0]
    for cell in table.columns[1].cells:
        cell.width = widths[1]
    for cell in table.columns[2].cells:
        cell.width = widths[2]
    for cell in table.columns[3].cells:
        cell.width = widths[3]
    for cell in table.columns[4].cells:
        cell.width = widths[4]

    # En-tête du tableau
    hdr_cells = table.rows[0].cells
    headers = ['DÉSIGNATION', 'QUANTITÉ', 'PRIX UNITAIRE', 'REMISE', 'TOTAL']
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        hdr_cells[idx].paragraphs[0].style.font.name = "Arial"
        hdr_cells[idx].paragraphs[0].style.font.size = Pt(10)
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[idx].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[idx].paragraphs[0].paragraph_format.space_after = Pt(0)
        hdr_cells[idx].paragraphs[0].runs[0].bold = True
        set_cell_background_color(hdr_cells[idx], colorBlueBackground)

    # Corps du tableau
    total_amount = 0
    for produit in produits:
        row_cells = table.add_row().cells
        row_cells[0].text = produit[2]
        row_cells[1].text = f'{produit[3]}' if produit[6] is None else f'{produit[3]} {produit[6]}'
        row_cells[2].text = f'{produit[4]:.2f} €'
        row_cells[3].text = f'{produit[5]:.2f} €'
        total = produit[3] * (produit[4] - produit[5])
        row_cells[4].text = f'{total:.2f} €'
        total_amount += total

        for cell in row_cells:
            cell.paragraphs[0].style.font.name = "Arial"
            cell.paragraphs[0].style.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Définir la hauteur des lignes
    table.rows[0].height = Inches(0.35433071)  # 0,9 cm
    for row in table.rows[1:]:
        row.height = Inches(0.23622047)  # 0,6 cm

    # Ajouter une bordure au tableau
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 2px
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), colorBlackBorder)
        tblBorders.append(border)
    tblPr.append(tblBorders)

    return total_amount

def add_footer_to_last_page(document):
    """Ajoute un pied de page au bas de la dernière page du document."""
    # Ajouter un espace avant le footer
    document.add_paragraph()

    # Premier paragraphe
    paragraph1 = document.add_paragraph()
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    text = "Ce présent "
    text_devis = "devis"
    text_rest = " est valable pendant 1 mois à compter de la date d'émission, si acceptation après, le montant sera révisé.\nPour le valider, renvoyez-le à l'adresse suivante avec la mention\n« BON POUR ACCORD » suivi de votre signature"

    run1 = paragraph1.add_run(text)
    run1.font.name = "Arial"
    run1.font.size = Pt(8)
    run1.font.color.rgb = colorGreyText

    run2 = paragraph1.add_run(text_devis)
    run2.font.name = "Arial"
    run2.font.size = Pt(8)
    run2.font.color.rgb = colorRed  # Rouge

    run3 = paragraph1.add_run(text_rest)
    run3.font.name = "Arial"
    run3.font.size = Pt(8)
    run3.font.color.rgb = colorGreyText

    # Deuxième paragraphe
    paragraph2 = document.add_paragraph()
    paragraph2.text = "NMG&CO EI - 1215 Route de la Fénerie 06580 Pégomas - Tél : 07.69.98.25.66 - nmgiovan@gmail.com\nN° SIRET : 843 241 399 000 16 - APE : 8121Z\nIdentité Bancaire : FR76 1460 7003 5569 4219 9717 134"
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = paragraph2.runs[0]
    run2.font.name = "Arial"
    run2.font.size = Pt(10)
    run2.font.color.rgb = colorBlackText

def save_document(document, dossier_id):
    """Ouvre une boîte de dialogue pour sauvegarder le document."""
    root = tk.Tk()
    root.withdraw()

    document_name = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Document Word", "*.docx")],
        initialfile=f'devis_{dossier_id}.docx',
        title="Enregistrer le devis sous..."
    )

    if document_name:
        try:
            document.save(document_name)
            print(f'Devis généré: {document_name}')
            return True
        except Exception as e:
            print(f"Erreur lors de la sauvegarde du devis: {e}")
            return False
    else:
        print("Enregistrement annulé")
        return False

def generate_devis(dossier_id):
    dossier = get_dossier(dossier_id)
    if not dossier:
        print(f"Erreur : Aucun dossier trouvé avec l'identifiant {dossier_id}.")
        return False

    produits = get_produits(dossier_id)
    document = Document()

    set_document_margins(document, top=0.5, bottom=0.5, left=0.5, right=0.5)  # Définit des marges de 1,27 cm (0,5 pouce)
    create_header(document)
    add_dossier_info(document, dossier)
    description_dossier(document, dossier)
    total_amount = add_produits_table(document, produits)

    document.add_paragraph(f"SOUS TOTAL : {total_amount:.2f} €")
    document.add_paragraph(f"TOTAL A REGLER : {total_amount:.2f} €")
    document.add_paragraph("SI ACCEPTATION DU DEVIS ACOMPTE A VERSER : 522.50 €\nTVA NON APPLICABLE - Art. 293 B du CGI")

    # Add footer content to the last page
    add_footer_to_last_page(document)

    return save_document(document, dossier_id)

def get_dossier(dossier_id):
    dossiers = get_dossiers()
    for dossier in dossiers:
        if dossier[0] == dossier_id:
            return dossier
    return None

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python generate_devis.py <dossier_id>")
        sys.exit(1)

    dossier_id = int(sys.argv[1])
    if not generate_devis(dossier_id):
        sys.exit(1)